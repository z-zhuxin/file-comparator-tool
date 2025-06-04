from docx import Document
from typing import List

def parse_word(file_path: str) -> str:
    """
    解析Word文档内容，返回纯文本
    
    参数:
        file_path: Word文档路径(.docx)
        
    返回:
        文档内容的纯文本字符串
    """
    try:
        doc = Document(file_path)
        text_parts = []
        
        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
                        
        return '\n'.join(text_parts)
        
    except Exception as e:
        raise ValueError(f"解析Word文档失败: {str(e)}")

def parse_word_to_lines(file_path: str) -> List[str]:
    """
    解析Word文档内容，返回按行分割的文本列表
    
    参数:
        file_path: Word文档路径(.docx)
        
    返回:
        文档内容的文本行列表
    """
    full_text = parse_word(file_path)
    return [line for line in full_text.split('\n') if line.strip()]